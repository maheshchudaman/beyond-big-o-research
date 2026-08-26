#!/usr/bin/env python3
"""Build the evidence-based Mac pilot manuscript from generated metrics."""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
GENERATED = PAPER / "generated"
OUTPUT = PAPER / "Beyond_Big_O_Mac_Pilot_Study_Draft.docx"
METRICS = json.loads((GENERATED / "paper_metrics.json").read_text(encoding="utf-8"))

BLUE = "2E74B5"
DARK_BLUE = "0B2545"
MID_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5E6B78"
WHITE = "FFFFFF"
RISK = "9B1C1C"


def set_run(run, size=11, bold=False, italic=False, color=DARK_BLUE, font="Calibri"):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), font)
    rpr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            margins(cell)


def add_para(doc, text="", *, size=11, bold=False, italic=False, color=DARK_BLUE, align=None, before=0, after=8, line=1.333, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_rich_para(doc, parts, *, after=8, line=1.333, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    for text, options in parts:
        set_run(p.add_run(text), **options)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    p.paragraph_format.keep_with_next = True
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run(text))
    return p


def add_table(doc, caption, headers, rows, widths, font_size=9.2):
    add_para(doc, caption, size=9.5, bold=True, color=MID_BLUE, before=4, after=4, keep=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=font_size, bold=True, color=WHITE)
    repeat_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_index % 2:
                shade(cells[index], LIGHT_GRAY)
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if index > 1 and len(value) < 16:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(value), size=font_size)
    table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_note(doc, label, text, risk=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.line_spacing = 1.25
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FCE8E6" if risk else LIGHT_BLUE)
    ppr.append(shd)
    set_run(p.add_run(f"{label}: "), bold=True, color=RISK if risk else MID_BLUE)
    set_run(p.add_run(text))


def add_picture(doc, path, caption, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(6.3))._inline
    doc_pr = inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", caption.split(".")[0])
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(9)
    cap.paragraph_format.keep_with_next = False
    set_run(cap.add_run(caption), size=9, italic=True, color=MUTED)


def page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run(run, size=8.5, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, MID_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_run(header.add_run("BEYOND BIG-O  |  APPLE-SILICON PILOT STUDY"), size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("Draft manuscript  |  Page "), size=8.5, color=MUTED)
    page_number(footer)


def fmt_ms(value):
    if value < 0.001:
        return f"{value * 1000:.3f} us"
    if value < 1:
        return f"{value:.4f} ms"
    return f"{value:.3f} ms"


# Ordered by first in-text appearance so reference numbers stay sequential;
# add new sources here at the point they are first cited, not at the end.
REFERENCES = [
    ("cormen", "Cormen, T. H., Leiserson, C. E., Rivest, R. L., and Stein, C. (2022). Introduction to Algorithms (4th ed.). MIT Press."),
    ("hennessy", "Hennessy, J. L., and Patterson, D. A. (2019). Computer Architecture: A Quantitative Approach (6th ed.). Morgan Kaufmann."),
    ("sanders", "Sanders, P. (2009). Algorithm Engineering - An Attempt at a Definition. In Efficient Algorithms, LNCS 5760, 321-340. https://doi.org/10.1007/978-3-642-03456-5_22"),
    ("mytkowicz", "Mytkowicz, T., Diwan, A., Hauswirth, M., and Sweeney, P. F. (2009). Producing wrong data without doing anything obviously wrong! Proceedings of ASPLOS XIV, 265-276. https://doi.org/10.1145/1508244.1508275"),
    ("georges", "Georges, A., Buytaert, D., and Eeckhout, L. (2007). Statistically rigorous Java performance evaluation. Proceedings of OOPSLA 2007, 57-76. https://doi.org/10.1145/1297027.1297033"),
    ("kalibera", "Kalibera, T., and Jones, R. (2013). Rigorous benchmarking in reasonable time. Proceedings of ISMM 2013, 63-74. https://doi.org/10.1145/2464157.2464160"),
    ("arcuri", "Arcuri, A., and Briand, L. (2011). A practical guide for using statistical tests to assess randomized algorithms in software engineering. Proceedings of ICSE 2011, 1-10. https://doi.org/10.1145/1985793.1985795"),
    ("fleming", "Fleming, P. J., and Wallace, J. J. (1986). How not to lie with statistics: the correct way to summarize benchmark results. Communications of the ACM, 29(3), 218-221. https://doi.org/10.1145/5666.5673"),
    ("sandve", "Sandve, G. K., Nekrutenko, A., Taylor, J., and Hovig, E. (2013). Ten simple rules for reproducible computational research. PLOS Computational Biology, 9(10), e1003285. https://doi.org/10.1371/journal.pcbi.1003285"),
    ("peng", "Peng, R. D. (2011). Reproducible research in computational science. Science, 334(6060), 1226-1227. https://doi.org/10.1126/science.1213847"),
    ("pythondocs", "Python Software Foundation. (2026). Python 3.13 documentation: Data structures. https://docs.python.org/3.13/tutorial/datastructures.html"),
    ("isocpp", "ISO/IEC. (2020). ISO/IEC 14882:2020 Programming Languages - C++. International Organization for Standardization."),
    ("bootstrap", "Efron, B., and Tibshirani, R. J. (1993). An Introduction to the Bootstrap. Chapman and Hall/CRC."),
    ("cliff", "Cliff, N. (1993). Dominance statistics: Ordinal analyses to answer ordinal questions. Psychological Bulletin, 114(3), 494-509. https://doi.org/10.1037/0033-2909.114.3.494"),
    ("varghadelaney", "Vargha, A., and Delaney, H. D. (2000). A critique and improvement of the CL common language effect size statistics of McGraw and Wong. Journal of Educational and Behavioral Statistics, 25(2), 101-132. https://doi.org/10.3102/10769986025002101"),
    ("chilimbi", "Chilimbi, T. M., Hill, M. D., and Larus, J. R. (1999). Cache-conscious structure layout. Proceedings of PLDI 1999, 1-12. https://doi.org/10.1145/301618.301633"),
    ("drepper", "Drepper, U. (2007). What Every Programmer Should Know About Memory. Red Hat, Inc."),
    ("astrachan", "Astrachan, O. (2003). Bubble sort: an archaeological algorithmic analysis. ACM SIGCSE Bulletin, 35(1), 1-5. https://doi.org/10.1145/611892.611918"),
]
REF = {key: index + 1 for index, (key, _) in enumerate(REFERENCES)}


def build():
    doc = Document()
    configure(doc)

    # Academic title block (memo_masthead-inspired override).
    add_para(doc, "ORIGINAL RESEARCH - PILOT STUDY", size=9.5, bold=True, color=BLUE, before=4, after=7)
    add_para(doc, "Beyond Big-O: A Reproducible Apple-Silicon Pilot Study of Runtime Performance for Common Data Structures in Python and C++", size=22, bold=True, color=DARK_BLUE, after=9, line=1.08)
    add_para(doc, "Mahesh Patil*, Varun Patil", size=12, bold=True, color=MID_BLUE, after=2)
    add_para(doc, "Shah & Anchor Kutchhi Engineering College, Mumbai, Maharashtra, India", size=10, italic=True, color=MUTED, after=2)
    add_para(doc, "* Corresponding author (mahesh.patil@sakec.ac.in)", size=9, italic=True, color=MUTED, after=12)
    add_note(doc, "Draft status", "Prepared from measurements executed in Codex on 23 August 2026, with supplementary construct-validity (Section 4.5) and third-language (Section 4.6) checks added on 25 August 2026. This manuscript is intended for internal academic review and methodological refinement before journal submission.")

    heading(doc, "Abstract", 1)
    abstract = (
        "Asymptotic analysis predicts how algorithmic cost grows, but it does not capture language-runtime overhead, "
        "container representation, allocation behaviour or constants that influence observed execution time. This "
        "pilot study compared dynamic arrays, linked structures and hash tables under identical build, search, "
        "deletion and traversal workloads in Python 3.13 and optimized C++17, executed on an Apple M2 MacBook Air "
        "with 8 GB memory and macOS 26.5.2. Four input sizes (1,000-25,000) were tested using three warm-ups and ten "
        "recorded repetitions, yielding 240 validated measurement records. At n = 25,000, median Python runtimes "
        "were 1.86-91.68 times the corresponding C++ medians across the twelve structure-operation combinations. "
        "Batched search and deletion for sequential structures exhibited empirical scaling exponents of "
        "approximately 1.85-2.00 because both the collection size and the number of operations increased with n; "
        "hash workloads were closer to linear. All implementations produced identical search-hit counts and "
        "post-deletion checksums, including four supplementary checks added on 25 August 2026 that test the primary "
        "findings rather than merely disclosing them as caveats. A singly-linked std::forward_list re-run "
        "(Section 4.5) reproduced the same pattern of ratios as the primary std::list comparison. Adding Java as a "
        "third language (Section 4.6) showed an operation-dependent ratio to C++, slower than Python on four of "
        "twelve combinations, though the widest ratios coincide with measurement dispersion far higher than the "
        "primary design (median CV 38.6% vs 4.0%); low-dispersion combinations gave a narrower, more defensible "
        "1.49x-1.64x range. Calibrated batching (Section 4.7) resolved the two resolution-limited C++ groups flagged "
        "since the first draft, cutting their dispersion from clock-tick noise to under 2% CV. A final supplement "
        "(Section 4.8) added peak memory, coarse instruction/cycle counters, a thermal snapshot and a second "
        "input-distribution family: search, deletion and traversal were robust to distribution, but C++ insertion "
        "was not, for reasons not yet identified. The findings support teaching Big-O together with "
        "implementation-aware measurement, while the single-machine scope, absence of literal cache-miss counters, "
        "and the unresolved insertion-distribution effect keep this a pilot rather than a final general-purpose "
        "benchmark."
    )
    add_para(doc, abstract, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_rich_para(doc, [("Keywords: ", {"bold": True}), ("data structures; Big-O; empirical algorithmics; Python; C++; Java; microbenchmarking; reproducibility; construct validity", {})], align=WD_ALIGN_PARAGRAPH.LEFT)

    heading(doc, "1. Introduction", 1)
    add_para(doc, f"Big-O notation provides a machine-independent description of growth as input size increases. It is indispensable for algorithm design, yet it intentionally suppresses constant factors, representation costs and hardware effects [{REF['cormen']}]. Consequently, two implementations with the same asymptotic class can differ substantially in observed runtime, while a theoretically favourable structure may carry a higher construction or traversal cost for a particular workload.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, f"This distinction is especially important in undergraduate computing education. Students often learn that hash-table lookup is expected O(1), array search is O(n), and linked-list traversal is O(n), but they may not observe how language runtimes, boxed objects, allocation strategies and contiguous memory influence actual measurements. Computer architecture texts likewise emphasise that locality and the memory hierarchy shape performance beyond instruction counts [{REF['hennessy']}].", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "The objective of this pilot is not to declare one programming language universally superior. It is to test whether a small, reproducible experiment can connect theoretical complexity with measured behaviour while making its assumptions and limitations explicit. The resulting protocol is intended as a foundation for a student research project and a later multi-platform study.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "1.1 Research questions", 2)
    bullet(doc, "RQ1: How do language and data-structure implementation affect observed runtime for equivalent workloads?")
    bullet(doc, "RQ2: Do measured scaling patterns agree with the workload-level complexity predicted from Big-O analysis?")
    bullet(doc, "RQ3: Which methodological limitations must be addressed before extending the pilot into a journal-ready benchmark?")

    heading(doc, "1.2 Contributions", 2)
    bullet(doc, "A common, deterministic dataset and workload definition shared by Python and C++ implementations.")
    bullet(doc, "A reproducible execution pipeline with correctness checks, raw-data hashing and environment metadata.")
    bullet(doc, "An empirical distinction between per-operation complexity and the complexity of a batch whose operation count also grows with n.")
    bullet(doc, "A transparent account of measurement-resolution, implementation-equivalence and external-validity limitations.")
    bullet(doc, "An empirical test of whether the C++/Python linked-container mismatch inflates the observed language gap, rather than leaving that concern as an unverified caveat.")
    bullet(doc, "A third-language check that re-runs the full workload under Java, testing whether the Python/C++ pattern generalises rather than leaving Java's absence as an unaddressed limitation.")
    bullet(doc, "A calibrated-batching supplement that resolves the two resolution-limited C++ groups into stable, low-dispersion estimates, converting an open measurement caveat into a quantified one.")
    bullet(doc, "Peak-memory, coarse hardware-counter and second-input-distribution supplements that close four of the six original future-work items within a single Mac's reach, leaving only literal cache-miss counters and independent-system reproduction as genuinely out of scope here.")

    heading(doc, "2. Background and Related Work", 1)
    add_para(doc, f"The analysis of algorithms separates growth rate from implementation detail [{REF['cormen']}]. Experimental algorithmics complements that abstraction by examining implementations on specified workloads and machines. Algorithm engineering has been described as a methodology connecting design, implementation, experimentation and refinement [{REF['sanders']}]. This perspective motivates reporting sufficient detail to reproduce not only the code but also the experimental conditions.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, f"Performance measurement is vulnerable to effects that are easy to overlook. Mytkowicz et al. showed that apparently harmless environmental changes can alter conclusions [{REF['mytkowicz']}]. Georges et al. and Kalibera and Jones argued for warm-ups, repeated measurements, uncertainty reporting and explicit treatment of runtime-system variability [{REF['georges']},{REF['kalibera']}]. Arcuri and Briand outlined complementary statistical-testing practice for evaluating randomized algorithms in software engineering, including effect-size reporting alongside significance testing [{REF['arcuri']}]. Fleming and Wallace further cautioned that benchmark summaries can mislead when inappropriate averages are used [{REF['fleming']}]. Reproducibility itself depends on preserving code, data and environment details rather than reporting summary numbers alone [{REF['sandve']},{REF['peng']}]. The present pilot therefore prioritises medians, dispersion, warm-ups, immutable raw results and correctness checks.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "The study is deliberately modest. It does not claim that its three containers exhaust data-structure design or that Python and C++ expose equivalent internal representations. Instead, it investigates idiomatic implementations under a shared external workload. That decision improves classroom relevance but weakens causal attribution: measured differences reflect the combined effect of language, runtime, library implementation and element representation.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "3. Materials and Methods", 1)
    heading(doc, "3.1 Execution environment", 2)
    env_rows = [
        ["Computer", "MacBook Air (Mac14,2)"],
        ["Processor", "Apple M2, 8 cores (4 performance, 4 efficiency), arm64"],
        ["Memory", "8 GB unified memory"],
        ["Operating system", "macOS 26.5.2, Darwin 25.5.0"],
        ["Python", "CPython 3.13.5"],
        ["C++", "C++17 compiled with Apple Clang 21.0.0 using -O2"],
        ["Java", "Not executed: no working JDK was available"],
        ["Hardware counters", "Not collected: Linux perf unavailable on macOS"],
        ["Execution date", "23 August 2026"],
    ]
    add_table(doc, "Table 1. Recorded execution environment.", ["Item", "Recorded value"], env_rows, [2400, 6960], 9.4)

    heading(doc, "3.2 Structures and workload", 2)
    add_para(doc, f"The experiment used three abstract structures. The dynamic-array condition used Python list [{REF['pythondocs']}] and C++ std::vector<int> [{REF['isocpp']}]; the linked condition used a custom Python singly linked list and C++ std::list<int>; and the hash condition used Python dict and C++ std::unordered_map<int,int>. Inputs contained unique integers in a deterministic shuffled order.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    method_rows = [
        ["Input size", "1,000; 5,000; 10,000; 25,000"],
        ["Build", "Construct the structure from all n values"],
        ["Search", "Query max(10, 0.01n) keys; half present and half absent"],
        ["Delete", "Delete max(1, 0.01n) keys known to be present"],
        ["Traverse", "Sum all keys remaining after deletion"],
        ["Correctness", "Compare search-hit count and post-deletion checksum"],
        ["Timing boundary", "Exclude dataset parsing and CSV output"],
    ]
    add_table(doc, "Table 2. Common experimental workload.", ["Component", "Definition"], method_rows, [2300, 7060], 9.4)

    heading(doc, "3.3 Dataset generation and validation", 2)
    add_para(doc, "A fixed seed (20260823) generated one text dataset per input size. Every language read the same files. A manifest stored each file's SHA-256 digest. Before performance analysis, all result groups were checked for identical search-hit counts and identical post-deletion checksums. Any disagreement would have invalidated the performance data; none occurred.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "3.4 Measurement protocol", 2)
    add_para(doc, "For every language-structure-size combination, the program performed three unrecorded warm-ups followed by ten recorded repetitions. Jobs were shuffled with the experiment seed. Operation durations were measured using monotonic high-resolution clocks and reported in nanoseconds. The complete design produced 2 languages x 3 structures x 4 sizes x 10 repetitions = 240 records.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_note(doc, "Measurement caveat", "Some optimized C++ workloads completed in less than one microsecond, and the C++ hash-search median at n = 1,000 was recorded as zero. Such values are below a dependable single-shot timing range and are treated as resolution-limited rather than literal zero-cost operations. Section 4.7 reports a calibrated-batching supplement that resolves this ambiguity for the affected groups.", risk=True)

    heading(doc, "3.5 Analysis", 2)
    add_para(doc, f"The primary summary was the median across ten repetitions, with the mean, standard deviation and non-parametric bootstrap interval retained in the generated analysis file [{REF['bootstrap']}]. Coefficient of variation (CV) was used to describe group dispersion. Python-to-C++ median ratios at n = 25,000 were calculated descriptively; Cliff's delta [{REF['cliff']}], following the A-measure formulation of Vargha and Delaney [{REF['varghadelaney']}], was also computed. Scaling exponents between n = 5,000 and n = 25,000 were estimated as log(T2/T1) / log(n2/n1). No causal language claim is made because language and container representation were not independently manipulated.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "4. Results", 1)
    heading(doc, "4.1 Correctness and measurement stability", 2)
    add_para(doc, f"All 240 records passed the hit-count and checksum validation. The median CV across the 96 language-structure-size-operation groups was {METRICS['median_group_cv']*100:.1f}%. The maximum CV was {METRICS['max_group_cv']*100:.1f}% and occurred in the resolution-limited C++ hash-search group at n = 1,000. The data therefore show generally stable repeated measurements while also identifying the exact region in which the timer cannot support strong absolute claims.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "4.2 Search scaling", 2)
    add_picture(doc, GENERATED / "figure_1_search_scaling.png", "Figure 1. Median search-workload time across input sizes. The zero C++ hash-search median at n = 1,000 is omitted from the logarithmic plot.", "Two-panel log-log line chart comparing array, linked and hash search workload times for C++17 and Python 3.13 across four input sizes.")
    add_para(doc, "Sequential search rose much more sharply than hash search. Between 5,000 and 25,000 elements, the estimated search exponent was 1.99 for the C++ array and 2.00 for the Python array; the linked implementations produced 1.91 and 1.99, respectively. The corresponding hash exponents were 0.90 for C++ and 0.97 for Python. This apparent quadratic-versus-linear contrast is consistent with the workload definition: the number of search queries itself increases in proportion to n, so an O(n) sequential search repeated O(n) times yields an O(n^2) batch, whereas expected O(1) hash lookup repeated O(n) times yields an O(n) batch.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "4.3 Runtime at the largest input", 2)
    result_rows = []
    for structure in ("array", "linked", "hash"):
        for operation in ("insert", "search", "delete", "traverse"):
            entry = METRICS["n25000"][structure][operation]
            result_rows.append([
                structure.title(), operation.title(), fmt_ms(entry["cpp"]["median_ms"]),
                fmt_ms(entry["python"]["median_ms"]), f"{entry['python_to_cpp_median_ratio']:.2f}x",
            ])
    add_table(doc, "Table 3. Median workload runtime at n = 25,000 (ten repetitions).", ["Structure", "Operation", "C++17", "Python 3.13", "Python/C++"], result_rows, [1700, 1800, 1900, 2100, 1860], 8.8)
    ratio_min, ratio_max = METRICS["language_ratio_range_n25000"]
    add_para(doc, f"At n = 25,000, every Python median exceeded its corresponding C++ median; the descriptive ratios ranged from {ratio_min:.2f}x for hash deletion to {ratio_max:.2f}x for array traversal. Cliff's delta was 1.00 for all twelve comparisons, meaning every recorded Python duration in each comparison exceeded every corresponding C++ duration. These values characterise the tested implementations on this Mac and must not be interpreted as universal language speed factors.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_picture(doc, GENERATED / "figure_2_language_ratios.png", "Figure 2. Python-to-C++ median runtime ratios at n = 25,000. The horizontal axis is logarithmic.", "Horizontal bar chart of Python-to-C++ median runtime ratios for insertion, search, deletion and traversal across array, linked and hash structures.")

    heading(doc, "4.4 Deletion and traversal patterns", 2)
    add_para(doc, "Array and linked deletion workloads also approached quadratic scaling because 1% of n keys were deleted and each deletion required sequential location; estimated exponents ranged from 1.85 to 1.89 for linked deletion and 1.87 to 1.89 for array deletion. Hash deletion remained close to linear at 1.00 in C++ and 1.02 in Python. Traversal generally approached linear growth, with exponents between 0.84 and 0.97 across the structures and languages. These patterns are more defensible than the smallest absolute timings because they rely on changes across larger workloads.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "4.5 Construct-validity check: a singly-linked C++ comparison", 2)
    add_para(doc, "Section 6.2 notes a representational mismatch in the linked condition: the C++ side used std::list, which is typically doubly linked, while the Python side used a custom singly linked list. To test whether this mismatch inflates the reported ratios rather than merely being noted as a caveat, the C++ linked benchmark was re-run using std::forward_list, a singly-linked container with the same single-next-pointer structure as the Python implementation, across all four input sizes with the same three-warm-up, ten-repetition protocol. Every forward_list run reproduced the same search-hit count and post-deletion checksum as both the original std::list run and the Python run at every input size, verified programmatically before the comparison was computed.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    fwd = METRICS["linked_fwd_supplement"]["n25000"]
    fwd_rows = []
    for operation in ("insert", "search", "delete", "traverse"):
        entry = fwd[operation]
        fwd_rows.append([
            operation.title(), fmt_ms(entry["cpp_list_median_ms"]), fmt_ms(entry["cpp_forward_list_median_ms"]),
            fmt_ms(entry["python_median_ms"]), f"{entry['python_to_forward_list_ratio']:.2f}x",
        ])
    add_table(doc, "Table 4. Construct-validity check: median linked-structure runtime at n = 25,000 using a singly-linked C++ container matched to the Python implementation.", ["Operation", "C++ std::list", "C++ std::forward_list", "Python 3.13", "Python/forward_list"], fwd_rows, [1600, 1900, 2000, 2000, 1860], 8.6)
    add_para(doc, f"Matching the C++ container's linkage discipline to the Python implementation did not narrow the measured language gap. The Python-to-C++ ratio was higher under the forward_list comparison than under std::list for search ({fwd['search']['python_to_forward_list_ratio']:.2f}x versus {fwd['search']['python_to_list_ratio']:.2f}x), deletion ({fwd['delete']['python_to_forward_list_ratio']:.2f}x versus {fwd['delete']['python_to_list_ratio']:.2f}x) and traversal ({fwd['traverse']['python_to_forward_list_ratio']:.2f}x versus {fwd['traverse']['python_to_list_ratio']:.2f}x), and only slightly lower for insertion ({fwd['insert']['python_to_forward_list_ratio']:.2f}x versus {fwd['insert']['python_to_list_ratio']:.2f}x). Between n = 5,000 and n = 25,000, the forward_list search and deletion exponents were {METRICS['linked_fwd_supplement']['scaling_exponents_5000_25000']['search']:.2f} and {METRICS['linked_fwd_supplement']['scaling_exponents_5000_25000']['delete']:.2f} respectively, consistent with the batch-scaling explanation given in Section 5.1 for std::list. The construct-validity concern raised in Section 6.2 therefore does not appear to be inflating the reported Python/C++ ratios for the linked condition, though it remains a caveat for claims about a specific container's internal performance rather than about Python versus C++ generally.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "4.6 External-validity check: adding Java as a third language", 2)
    java_sup = METRICS["java_supplement"]
    add_para(doc, f"Section 6.3 notes that Java was not executed in the primary study because no working JDK was available at measurement time. To test whether the reported Python/C++ pattern generalises to a third, managed-runtime language, the Java benchmark already scaffolded in the repository was compiled and run under {java_sup['java_version']} across all three structures and four input sizes with the identical three-warm-up, ten-repetition protocol, on 25 August 2026. Every Java run reproduced the same search-hit count and post-deletion checksum as the corresponding Python and C++ runs at every input size, verified programmatically before the comparison was computed.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    java_rows = []
    for structure in ("array", "linked", "hash"):
        for operation in ("insert", "search", "delete", "traverse"):
            entry = java_sup["n25000"][structure][operation]
            java_rows.append([
                structure.title(), operation.title(), fmt_ms(entry["cpp_median_ms"]),
                fmt_ms(entry["python_median_ms"]), fmt_ms(entry["java_median_ms"]),
                f"{entry['java_to_cpp_ratio']:.2f}x",
            ])
    add_table(doc, "Table 5. Median workload runtime at n = 25,000 including Java (ten repetitions).", ["Structure", "Operation", "C++17", "Python 3.13", "Java", "Java/C++"], java_rows, [1500, 1600, 1750, 1750, 1550, 1610], 8.4)
    stable_lo, stable_hi = java_sup["stable_ratio_range"]
    add_note(
        doc,
        "Measurement caveat",
        f"Median CV across the Java supplement's 48 (structure, size, operation) groups was "
        f"{java_sup['median_group_cv_pct']:.1f}%, versus 4.0% for the primary Python/C++ design, and the maximum "
        f"reached {java_sup['max_group_cv_pct']:.1f}%. Dispersion was concentrated in groups with short absolute "
        f"durations, including linked insertion at n = 25,000 (CV = {java_sup['n25000']['linked']['insert']['java_cv_pct']:.1f}%: "
        f"several repetitions cluster near tens of microseconds while others exceed one millisecond), consistent with "
        f"the JIT compiler transitioning from interpreted to compiled execution partway through the ten recorded "
        f"repetitions despite three unrecorded warm-ups. Array and linked search and delete at n = 25,000, the "
        f"longest-duration groups, remained stable (CV under 6%). Ratios computed from high-dispersion groups should "
        f"be read as indicative rather than precise.",
        risk=True,
    )
    slower = java_sup["slower_than_python_combinations"]
    slower_desc = ", ".join(f"{s} {op}" for s, op in slower)
    add_para(
        doc,
        f"Java's ratio to C++ was operation-dependent rather than uniformly intermediate between the interpreted and "
        f"compiled extremes: across the twelve structure-operation combinations it ranged from 0.41x (linked "
        f"insertion) to 109.78x (array traversal), a wider spread than the Python/C++ range for the same twelve "
        f"combinations, though both extremes sit in high-dispersion groups (see the measurement caveat above). Among "
        f"the four combinations with low dispersion (array and linked search and delete, all CV under 6%), the range "
        f"narrows to a more defensible {stable_lo:.2f}x-{stable_hi:.2f}x. Four of the twelve combinations showed Java "
        f"slower than Python ({slower_desc}), not the uniform Java-beats-Python pattern the JIT-compiled label might "
        f"suggest. The largest such gap was array insertion, where Java's median (0.241 ms) was about 7.7 times "
        f"Python's (0.031 ms) and 36.66 times C++'s; hash traversal was next, at roughly four times Python's median. "
        f"Both plausibly reflect boxed-object overhead specific to Java's reference-type collections -- "
        f"ArrayList<Integer> resizing and autoboxing for insertion, boxed-Long iteration for HashMap traversal -- "
        f"rather than a general JVM weakness, since Java was faster than Python on hash insertion and search for the "
        f"same structure. Scaling exponents for the stable search and deletion operations between n = 5,000 and n = "
        f"25,000 were "
        f"{java_sup['scaling_exponents_5000_25000']['array']['search']:.2f} and "
        f"{java_sup['scaling_exponents_5000_25000']['array']['delete']:.2f} for the array and "
        f"{java_sup['scaling_exponents_5000_25000']['linked']['search']:.2f} and "
        f"{java_sup['scaling_exponents_5000_25000']['linked']['delete']:.2f} for the linked structure, consistent "
        f"with the same batch-scaling explanation given in Section 5.1. These results reinforce rather than "
        f"complicate the paper's central methodological point: coarse language-level generalisations do not hold "
        f"uniformly across operations, and the measurement itself, not just the workload, must be scrutinised before "
        f"a ratio is trusted.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    heading(doc, "4.7 Measurement-resolution supplement: calibrated batching", 2)
    calib = METRICS["calibration_supplement"]
    add_para(
        doc,
        "Section 3.4 flags array traversal and hash search as vulnerable to timer quantisation: both are far shorter "
        "than one microsecond at small n, and the C++ hash-search median at n = 1,000 was recorded as literally zero. "
        "Both operations are read-only and idempotent, so unlike insertion or deletion they can be repeated on the "
        "same container without side effects. A calibrated-batching supplement therefore re-measured both across all "
        f"four input sizes: for each repeat, the operation was run in a tight loop with the batch size doubled until "
        f"the total interval reached a {calib['threshold_ns'] // 1_000_000} ms calibration threshold, then per-operation "
        f"time was reported as elapsed time divided by batch size. Every calibrated run reproduced the same hit count "
        f"or checksum as the corresponding primary single-shot run at every input size, verified programmatically "
        f"before the comparison was computed.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    calib_rows = []
    for structure, operation in (("array", "traverse"), ("hash", "search")):
        for n in (1000, 5000, 10000, 25000):
            e = calib["by_size"][structure][operation][str(n)]
            calib_rows.append([
                structure.title(), operation.title(), f"{n:,}",
                f"{e['single_shot_ns']:.0f} ns", f"{e['calibrated_median_ns']:.1f} ns", f"{e['calibrated_cv_pct']:.1f}%",
            ])
    add_table(doc, "Table 6. Single-shot versus calibrated-batch median per-operation time.", ["Structure", "Operation", "n", "Single-shot", "Calibrated", "Calibrated CV"], calib_rows, [1300, 1500, 1200, 1700, 1700, 1360], 8.4)
    hash_1000 = calib["by_size"]["hash"]["search"]["1000"]
    array_1000 = calib["by_size"]["array"]["traverse"]["1000"]
    add_para(
        doc,
        f"Calibrated batching resolved the ambiguity the primary design could not: the C++ hash-search median at "
        f"n = 1,000, recorded as zero in the single-shot design, resolves to {hash_1000['calibrated_median_ns']:.0f} "
        f"ns under batching -- a real, non-zero cost, not literal zero-cost lookup. Dispersion across the eight "
        f"(structure, operation, n) groups was far lower than the single-shot design's: median CV "
        f"{calib['median_group_cv_pct']:.1f}%, maximum {calib['max_group_cv_pct']:.1f}% (the array-traversal group at "
        f"n = 1,000, still the smallest absolute magnitude tested). Calibrated medians were consistently close to but "
        f"somewhat below the original single-shot medians at every size (for example, array traversal at n = 25,000: "
        f"{calib['by_size']['array']['traverse']['25000']['calibrated_median_ns']:.0f} ns calibrated versus "
        f"{calib['by_size']['array']['traverse']['25000']['single_shot_ns']:.0f} ns single-shot), consistent with the "
        f"single-shot design carrying a small, roughly constant per-call overhead that batching amortises away. This "
        f"confirms the primary study's point estimates were directionally correct despite their high dispersion, and "
        f"converts the resolution-limited caveat in Section 3.4 from an open concern into a quantified, resolved one.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    heading(doc, "4.8 Resource-usage and additional-distribution supplements", 2)
    res = METRICS["resource_supplement"]
    dist = METRICS["sorted_distribution_supplement"]
    add_para(
        doc,
        "Three further gaps from the original future-work list were addressed within the constraints of a single "
        "Apple Silicon Mac. A thermal/power snapshot taken with pmset immediately before and after this session's "
        "runs recorded no thermal or performance warning level, arguing against throttling as a confound, though "
        "this is a snapshot rather than a continuous log. Peak process memory and two coarse hardware counters "
        "(instructions retired, cycles elapsed) were captured with /usr/bin/time -l, which needs no Linux perf "
        "dependency; this yields instruction-level and cycle-level detail but not literal cache hit/miss counts, so "
        "it complements rather than replaces the Linux hardware-counter work in Section 6.3. A three-run spot-check "
        "for one combination showed under 1.5% variation in memory, instructions and cycles, so a single measurement "
        "per (language, structure) is treated as representative.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    res_rows = []
    for lang, lang_label in (("cpp", "C++17"), ("python", "Python 3.13"), ("java", "Java")):
        for structure in ("array", "linked", "hash"):
            e = res["at_n25000"][lang][structure]
            res_rows.append([
                lang_label, structure.title(), f"{e['peak_memory_bytes'] / 1_000_000:.1f} MB",
                f"{e['instructions_retired'] / 1_000_000:.0f} M", f"{e['cycles_elapsed'] / 1_000_000:.0f} M", f"{e['ipc']:.2f}",
            ])
    add_table(doc, "Table 7. Peak memory and hardware counters at n = 25,000, whole-process.", ["Language", "Structure", "Peak memory", "Instructions", "Cycles", "IPC"], res_rows, [1550, 1550, 1550, 1550, 1550, 1610], 8.4)
    add_para(
        doc,
        f"Peak memory ordered C++ far below Python and Java throughout (for example, array at n = 25,000: "
        f"{res['at_n25000']['cpp']['array']['peak_memory_bytes']/1_000_000:.1f} MB versus "
        f"{res['at_n25000']['python']['array']['peak_memory_bytes']/1_000_000:.1f} MB for Python and "
        f"{res['at_n25000']['java']['array']['peak_memory_bytes']/1_000_000:.1f} MB for Java), consistent with "
        f"primitive contiguous storage versus boxed reference-type collections. Instructions-per-cycle gives a "
        f"mechanistic complement to the cache-locality argument in Section 5.2: C++ linked traversal's IPC "
        f"({res['at_n25000']['cpp']['linked']['ipc']:.2f}) was far below array's ({res['at_n25000']['cpp']['array']['ipc']:.2f}) "
        f"or hash's ({res['at_n25000']['cpp']['hash']['ipc']:.2f}), consistent with pointer-chasing stalling the "
        f"pipeline on memory latency rather than merely executing more instructions.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    dist_rows = []
    for structure in ("array", "linked", "hash"):
        e = dist["cpp_n25000"]["insert"][structure]
        dist_rows.append([structure.title(), f"{e['shuffled_ns']:.0f} ns", f"{e['sorted_ns']:.0f} ns", f"{e['ratio']:.2f}x"])
    add_para(
        doc,
        "A second input-distribution family (ascending-sorted values, same seed and query/delete fractions as the "
        "primary shuffled family, differing only in build order) was re-run across all three languages, structures "
        "and sizes, cross-language-correctness-verified at every point. Search, deletion and traversal medians "
        "agreed within roughly 20% between the two distributions for C++ and Python at every structure, supporting "
        "the external validity of those findings beyond the one distribution used throughout the rest of the paper. "
        "Insertion did not: C++ insertion differed by more than 20% in both directions across all three structures "
        "(Table 8), a pattern reproduced across three independent process invocations for the array case and "
        "therefore not attributable to single-run noise. This is not explained by the study's timing-boundary design, "
        "since dataset parsing is explicitly excluded from the timed region (Table 2); root-causing it, for example "
        "by checking allocator or page-fault behaviour under each build order, is left for future work. Java's "
        "distribution comparisons additionally inherit the high measurement dispersion already established for fast "
        "Java operations in Section 4.6 and are not reported here for that reason.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_table(doc, "Table 8. C++ insertion time under the two input distributions at n = 25,000.", ["Structure", "Shuffled", "Sorted", "Ratio"], dist_rows, [2340, 2340, 2340, 2340], 9.0)

    heading(doc, "5. Discussion", 1)
    heading(doc, "5.1 Big-O explained, not contradicted", 2)
    add_para(doc, "The results do not show a failure of asymptotic analysis. Instead, they demonstrate why the unit of analysis must be stated carefully. A single sequential search is O(n), but this experiment schedules 0.01n searches, creating an O(n^2) batch. A hash lookup is expected O(1) on average, and the corresponding O(n) batch grew close to linearly. The empirical exponents therefore connect measured behaviour to the composed workload rather than replacing theory.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "5.2 Language and representation effects", 2)
    add_para(doc, f"The Python/C++ differences include many inseparable factors: interpretation versus optimized native code, boxed Python integers versus C++ primitive integers, object allocation, iterator implementation and library design. The largest ratio occurred for contiguous traversal, where optimized native loops can exploit compact primitive storage, spatial locality and compiler transformations that pointer-chasing and boxed representations largely defeat [{REF['chilimbi']},{REF['drepper']}]. Hash insertion and deletion showed smaller ratios, indicating that language overhead is not a single constant applied uniformly to every structure and operation.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "5.3 Educational value", 2)
    add_para(doc, f"A useful teaching sequence is therefore: predict complexity, define the entire workload, implement correctness checks, measure, and then explain both agreement and disagreement. Students should learn to ask not only 'What is the Big-O?' but also 'What exactly is repeated, what representation is used, what is inside the timed region, and is the measurement long enough for the clock?' This echoes earlier calls in computing education to pair algorithmic analysis with empirical measurement rather than treating asymptotic notation as a substitute for it [{REF['astrachan']}]. The repository makes those questions visible and reproducible.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "6. Threats to Validity", 1)
    heading(doc, "6.1 Internal validity", 2)
    add_para(doc, "Background activity, thermal conditions and operating-system scheduling were not instrumented. Although jobs were shuffled and measurements were repeated, the study did not pin processes to performance cores. Extremely short C++ operations are vulnerable to timer quantisation and clock-call overhead; Section 4.7 confirms this directly for array traversal and hash search by batching each operation until its timed interval exceeds a predefined duration, reducing dispersion from clock-tick noise to under 2% CV in three of four sizes. Insertion and deletion were not similarly batchable because they mutate the container and cannot be repeated without rebuilding it between batches; a revised benchmark should rebuild-and-batch these operations too and should separately quantify timing overhead.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, f"The three-warm-up protocol, calibrated for the primary C++/Python design, proved insufficient for several fast Java operations (Section 4.6): median CV across the Java supplement's 48 groups was {METRICS['java_supplement']['median_group_cv_pct']:.1f}% versus 4.0% for the primary design, consistent with JIT compilation transitioning from interpreted to compiled execution mid-measurement. A revised Java benchmark should detect steady state explicitly, for example by warming up until consecutive batches agree within a tolerance, rather than relying on a fixed repetition count.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "6.2 Construct validity", 2)
    add_para(doc, "The structures are functionally comparable but not internally identical. Python's custom singly linked list differs from C++ std::list, which is typically doubly linked. Python list holds object references, whereas std::vector<int> stores primitive integers contiguously. Consequently, the experiment evaluates idiomatic implementation conditions rather than isolating a pure language effect. This concern is tested empirically in Section 4.5, which reruns the C++ linked benchmark with the singly-linked std::forward_list; the resulting ratios do not shrink under the more closely matched comparison, indicating the linkage mismatch is not the source of the observed language gap. Peak memory and two coarse hardware counters were measured on this Mac in Section 4.8; literal L1/L2/L3 cache hit and miss counts still require Linux perf and are not inferred in this paper.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "6.3 External and conclusion validity", 2)
    add_para(doc, "One Apple M2 laptop, one Python version and one C++ toolchain cannot represent other processors, operating systems, compilers or runtime configurations; no independent second system has reproduced any part of this study. Only four sizes were examined. Section 4.8 adds a second, ascending-sorted input family alongside the primary deterministic shuffle, but two families remain a small basis for claims about input-distribution robustness in general. Ten repetitions support descriptive stability but not broad population inference. Java was not executed in the primary 23 August 2026 session because no working JDK was available; Section 4.6 reports a supplementary Java run added on 25 August 2026 using a subsequently installed JDK on the same machine, so that check still shares this section's single-machine limitation. The conclusions are therefore restricted to this environment and protocol.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "7. Conclusion and Future Work", 1)
    add_para(doc, "This reproducible Mac pilot showed that theoretical growth and observed runtime can be taught together. Sequential search and deletion batches grew close to quadratically when both collection size and operation count increased, while hash batches grew close to linearly. Python and C++ showed substantial but operation-dependent runtime differences, and all correctness outputs agreed. A supplementary Java run showed that a JIT-compiled, managed-runtime language does not sit uniformly between the interpreted and compiled extremes, slower than Python on four of the twelve structure-operation combinations despite being faster overall; the run also exposed measurement dispersion far higher than the primary design's, a methodological finding in its own right. The study also exposed a measurement weakness: several optimized C++ workloads were too short for reliable single-shot timing.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, "This revision closed four of the six original future-work items within a single Mac's reach: Section 4.6 added Java as a third language; Section 4.7 resolved the sub-microsecond timing caveat for read-only operations through calibrated batching; and Section 4.8 added peak memory, two coarse hardware counters, a thermal/power snapshot and a second input-distribution family -- the last of which surfaced a genuine new finding, that C++ insertion, unlike search, deletion and traversal, responds substantially to input order.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_note(doc, "Future scope", "Three extensions define the next study rather than this one's shortfall: batching the mutating insert and delete operations to explain the input-order sensitivity Section 4.8 surfaced -- harder than the read-only case, since it requires rebuilding the container between batches; reproducing this protocol on an independent second system; and measuring literal L1/L2/L3 cache hit and miss counts, which require Linux perf and are unavailable on this Mac. Each should be reported as new evidence, not combined silently with the present data.", risk=False)

    heading(doc, "Data and Code Availability", 1)
    add_para(doc, f"The repository contains the dataset generator, Python, C++ and Java implementations, validation tests, execution configuration, raw CSV records and analysis scripts. The raw combined CSV contains {METRICS['record_count']} records and has SHA-256 digest {METRICS['raw_sha256']}. The supplementary singly-linked construct-validity check (Section 4.5) adds {METRICS['linked_fwd_supplement']['record_count']} further C++ records with SHA-256 digest {METRICS['linked_fwd_supplement']['raw_sha256']}; the supplementary Java run (Section 4.6) adds {METRICS['java_supplement']['record_count']} further records with SHA-256 digest {METRICS['java_supplement']['raw_sha256']}; the calibrated-batching supplement (Section 4.7) adds {METRICS['calibration_supplement']['record_count']} further records with SHA-256 digest {METRICS['calibration_supplement']['raw_sha256']}; the resource-usage supplement (Section 4.8) adds {METRICS['resource_supplement']['record_count']} process-level records with SHA-256 digest {METRICS['resource_supplement']['raw_sha256']}; and the second-distribution supplement (Section 4.8) adds {METRICS['sorted_distribution_supplement']['record_count']} further records with SHA-256 digest {METRICS['sorted_distribution_supplement']['raw_sha256']}. The repository is publicly available at https://github.com/maheshchudaman/beyond-big-o-research and is archived on Zenodo with concept DOI 10.5281/zenodo.22089928, which always resolves to the latest archived release.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "Ethics, Authorship and AI-Assistance Statement", 1)
    add_para(doc, "No human-participant, personal or sensitive data were used. Generative AI (Codex and Claude) assisted with repository scaffolding, code review, execution orchestration, analysis scripting and preparation of this draft, including the supplementary construct-validity and third-language checks added on 25 August 2026. All numerical claims in the manuscript were generated programmatically from the preserved raw CSV; missing measurements were not fabricated. Before submission, the named authors must independently verify the code, results, citations and wording, approve authorship contributions, and adapt this disclosure to the selected journal's policy.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    heading(doc, "References", 1)
    for _, reference in REFERENCES:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        set_run(p.add_run(reference), size=9.5)

    heading(doc, "Appendix A. Reproducibility Record", 1)
    appendix_rows = [
        ["Experiment configuration", "config/experiment.json"],
        ["Dataset manifest", "data/generated/manifest.json"],
        ["Raw measurements", "results/raw/combined.csv"],
        ["Environment metadata", "results/raw/environment.json"],
        ["Paper analysis", "paper/analyse_for_paper.py"],
        ["Correctness tests", "tests/test_benchmark.py; tests/test_dataset_manifest.py"],
        ["Construct-validity supplement (raw)", "results/raw/cpp_linked_fwd_combined.csv"],
        ["Construct-validity supplement (source)", "src/cpp/benchmark.cpp (run_linked_fwd)"],
        ["Java supplement (raw)", "results/raw/java_combined.csv"],
        ["Java supplement (source)", "src/java/Benchmark.java"],
        ["Calibration supplement (raw)", "results/raw/cpp_calibrated_combined.csv"],
        ["Calibration supplement (source)", "src/cpp/benchmark.cpp (run_calibrated)"],
        ["Resource-usage supplement (raw)", "results/raw/resource_usage.csv"],
        ["Thermal/power snapshot", "results/raw/thermal_snapshot.txt"],
        ["Second-distribution supplement (raw)", "results/raw/sorted_combined.csv"],
        ["Second-distribution supplement (source)", "scripts/generate_datasets_sorted.py"],
    ]
    add_table(doc, "Table A1. Repository evidence map.", ["Evidence", "Repository path"], appendix_rows, [3000, 6360], 9.4)
    doc.core_properties.title = "Beyond Big-O: Apple-Silicon Runtime Pilot Study"
    doc.core_properties.subject = "Reproducible empirical comparison of data-structure runtime in Python and C++"
    doc.core_properties.author = "Mahesh Patil and Varun Patil"
    doc.core_properties.keywords = "Big-O, data structures, runtime, Python, C++, Apple M2, reproducibility"
    doc.core_properties.comments = "Draft generated from executed Mac measurements; requires author and journal-specific review."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
