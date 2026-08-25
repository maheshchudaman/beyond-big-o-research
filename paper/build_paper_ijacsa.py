#!/usr/bin/env python3
"""Build an IJACSA-template-conformant version of the Beyond Big-O manuscript.

Reuses the numbers already validated in paper/generated/paper_metrics.json
(the same source of truth as build_paper.py) but reflows the content into
the official SAI/IJACSA two-column Word template (paper/ijacsa/SAI_PAPER_FORMAT.docx,
downloaded from https://thesai.org/DownloadFileHandler.ashx?filen=SAI_PAPER_FORMAT),
using that template's own named styles and section/column geometry so the
output actually conforms to the journal's required format rather than just
looking similar to it.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
GENERATED = PAPER / "generated"
TEMPLATE = PAPER / "ijacsa" / "SAI_PAPER_FORMAT.docx"
OUTPUT = PAPER / "ijacsa" / "Beyond_Big_O_IJACSA_Ready.docx"
METRICS = json.loads((GENERATED / "paper_metrics.json").read_text(encoding="utf-8"))

# Geometry read directly out of the official template's sectPr blocks.
ONECOL_MARGIN = Twips(893)
TWOCOL_MARGIN = Twips(907)


# Ordered by first in-text appearance, same citation keys/order already
# validated in build_paper.py -- only the rendered string changes (IEEE
# style here instead of APA), the [n] numbers stay identical.
REFERENCES = [
    ("cormen", "T. H. Cormen, C. E. Leiserson, R. L. Rivest, and C. Stein, Introduction to Algorithms, 4th ed. Cambridge, MA, USA: MIT Press, 2022."),
    ("hennessy", "J. L. Hennessy and D. A. Patterson, Computer Architecture: A Quantitative Approach, 6th ed. Cambridge, MA, USA: Morgan Kaufmann, 2019."),
    ("sanders", "P. Sanders, “Algorithm engineering – an attempt at a definition,” in Efficient Algorithms, LNCS 5760, Berlin, Germany: Springer, 2009, pp. 321–340, doi: 10.1007/978-3-642-03456-5_22."),
    ("mytkowicz", "T. Mytkowicz, A. Diwan, M. Hauswirth, and P. F. Sweeney, “Producing wrong data without doing anything obviously wrong!,” in Proc. ASPLOS XIV, 2009, pp. 265–276, doi: 10.1145/1508244.1508275."),
    ("georges", "A. Georges, D. Buytaert, and L. Eeckhout, “Statistically rigorous Java performance evaluation,” in Proc. OOPSLA 2007, 2007, pp. 57–76, doi: 10.1145/1297027.1297033."),
    ("kalibera", "T. Kalibera and R. Jones, “Rigorous benchmarking in reasonable time,” in Proc. ISMM 2013, 2013, pp. 63–74, doi: 10.1145/2464157.2464160."),
    ("arcuri", "A. Arcuri and L. Briand, “A practical guide for using statistical tests to assess randomized algorithms in software engineering,” in Proc. ICSE 2011, 2011, pp. 1–10, doi: 10.1145/1985793.1985795."),
    ("fleming", "P. J. Fleming and J. J. Wallace, “How not to lie with statistics: the correct way to summarize benchmark results,” Commun. ACM, vol. 29, no. 3, pp. 218–221, 1986, doi: 10.1145/5666.5673."),
    ("sandve", "G. K. Sandve, A. Nekrutenko, J. Taylor, and E. Hovig, “Ten simple rules for reproducible computational research,” PLOS Comput. Biol., vol. 9, no. 10, e1003285, 2013, doi: 10.1371/journal.pcbi.1003285."),
    ("peng", "R. D. Peng, “Reproducible research in computational science,” Science, vol. 334, no. 6060, pp. 1226–1227, 2011, doi: 10.1126/science.1213847."),
    ("pythondocs", "Python Software Foundation, “Python 3.13 documentation: Data structures,” 2026. [Online]. Available: https://docs.python.org/3.13/tutorial/datastructures.html"),
    ("isocpp", "ISO/IEC, ISO/IEC 14882:2020 Programming Languages – C++. Geneva, Switzerland: International Organization for Standardization, 2020."),
    ("bootstrap", "B. Efron and R. J. Tibshirani, An Introduction to the Bootstrap. Boca Raton, FL, USA: Chapman and Hall/CRC, 1993."),
    ("cliff", "N. Cliff, “Dominance statistics: Ordinal analyses to answer ordinal questions,” Psychol. Bull., vol. 114, no. 3, pp. 494–509, 1993, doi: 10.1037/0033-2909.114.3.494."),
    ("varghadelaney", "A. Vargha and H. D. Delaney, “A critique and improvement of the CL common language effect size statistics of McGraw and Wong,” J. Educ. Behav. Stat., vol. 25, no. 2, pp. 101–132, 2000, doi: 10.3102/10769986025002101."),
    ("chilimbi", "T. M. Chilimbi, M. D. Hill, and J. R. Larus, “Cache-conscious structure layout,” in Proc. PLDI 1999, 1999, pp. 1–12, doi: 10.1145/301618.301633."),
    ("drepper", "U. Drepper, “What every programmer should know about memory,” Red Hat, Inc., 2007."),
    ("astrachan", "O. Astrachan, “Bubble sort: an archaeological algorithmic analysis,” ACM SIGCSE Bull., vol. 35, no. 1, pp. 1–5, 2003, doi: 10.1145/611892.611918."),
]
REF = {key: index + 1 for index, (key, _) in enumerate(REFERENCES)}


def para(doc, text="", style=None, align=None, space_after=None):
    p = doc.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p


def corresponding_author_note(doc, text):
    # The template's "sponsors" style carries a floating w:framePr + top border
    # (it was designed for a fixed sponsor-logo box) — using it here detaches
    # this line from normal flow in real Word, even though it looks fine in
    # LibreOffice. Build a plain centered paragraph instead.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    return p


def set_section(section, *, cols=1, margin=TWOCOL_MARGIN, header=None, footer=None):
    section.left_margin = margin
    section.right_margin = margin
    if header is not None:
        section.header_distance = header
    if footer is not None:
        section.footer_distance = footer
    sectPr = section._sectPr
    for old in sectPr.findall(qn("w:cols")):
        sectPr.remove(old)
    cols_el = OxmlElement("w:cols")
    if cols == 2:
        cols_el.set(qn("w:num"), "2")
        cols_el.set(qn("w:space"), "360")
    else:
        cols_el.set(qn("w:space"), "720")
    # w:cols must precede docGrid per schema order; insert before docGrid if present.
    docgrid = sectPr.find(qn("w:docGrid"))
    if docgrid is not None:
        docgrid.addprevious(cols_el)
    else:
        sectPr.append(cols_el)


def new_section(doc, *, cols, header=None, footer=None):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_section(section, cols=cols, header=header, footer=footer)
    return section


def add_table(doc, caption, headers, rows, col_widths_in):
    # "14_Table Head" auto-numbers via the template's own list definition (numId=9).
    para(doc, caption, style="14_Table Head")
    table = doc.add_table(rows=1, cols=len(headers))
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(header)
        cell.paragraphs[0].style = doc.styles["table col head"]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].paragraphs[0].text = ""
            cells[index].paragraphs[0].add_run(value)
            cells[index].paragraphs[0].style = doc.styles["table copy"]
            if index > 0:
                cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        for index, width in enumerate(col_widths_in):
            row.cells[index].width = Inches(width)
    for col_index, width in enumerate(col_widths_in):
        table.columns[col_index].width = Inches(width)
    para(doc, "", space_after=Pt(6))
    return table


def add_figure(doc, path, caption, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    # "16_Figure Caption" auto-numbers via the template's own list definition (numId=2).
    cap = para(doc, caption, style="16_Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fmt_ms(value):
    if value < 0.001:
        return f"{value * 1000:.3f} us"
    if value < 1:
        return f"{value:.4f} ms"
    return f"{value:.3f} ms"


def build():
    doc = Document(str(TEMPLATE))
    body = doc.element.body
    for child in list(body):
        body.remove(child)
    # Re-create a minimal trailing sectPr so the document stays valid; this
    # becomes doc.sections[0], which we then treat as the Title section.
    sect_pr = OxmlElement("w:sectPr")
    for tag, attrs in (
        ("w:pgSz", {"w:w": "12240", "w:h": "15840", "w:code": "1"}),
        ("w:pgMar", {"w:top": "1080", "w:right": "893", "w:bottom": "1440", "w:left": "893", "w:header": "720", "w:footer": "720", "w:gutter": "0"}),
        ("w:cols", {"w:space": "720"}),
        ("w:docGrid", {"w:linePitch": "360"}),
    ):
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        sect_pr.append(el)
    body.append(sect_pr)

    # ---- Title section (single column, full width) ----
    para(doc, "Beyond Big-O: A Reproducible Apple-Silicon Pilot Study of Runtime Performance for Common Data Structures in Python and C++", style="01_Paper Title")

    new_section(doc, cols=1, header=Twips(720), footer=Twips(720))
    para(doc, "Mahesh Patil¹*, Varun Patil²", style="Author")
    para(doc, "Shah & Anchor Kutchhi Engineering College, Mumbai, Maharashtra, India¹", style="Affiliation")
    para(doc, "Shah & Anchor Kutchhi Engineering College, Mumbai, Maharashtra, India²", style="Affiliation")
    corresponding_author_note(doc, "* Corresponding author (mahesh.patil@sakec.ac.in)")

    # ---- Main two-column body ----
    new_section(doc, cols=2)
    abstract = (
        "Abstract—Asymptotic analysis predicts how algorithmic cost grows, but it does not capture language-runtime "
        "overhead, container representation, allocation behaviour or constants that influence observed execution time. "
        "This pilot study compared dynamic arrays, linked structures and hash tables under identical build, search, "
        "deletion and traversal workloads in Python 3.13 and optimized C++17. Experiments were executed on an Apple M2 "
        "MacBook Air with 8 GB memory and macOS 26.5.2. Four input sizes (1,000-25,000) were tested using three "
        "warm-ups and ten recorded repetitions, yielding 240 validated measurement records. At n = 25,000, median "
        "Python runtimes were 1.86-91.68 times the corresponding C++ medians across the twelve structure-operation "
        "combinations; the upper end of this range reflects a submicrosecond, resolution-limited C++ traversal median "
        "rather than a stable absolute figure. Batched search and deletion for sequential structures exhibited "
        "empirical scaling exponents of approximately 1.85-2.00 because both the collection size and the number of "
        "operations increased with n; hash workloads were closer to linear. All implementations produced identical "
        "search-hit counts and post-deletion checksums, including a supplementary re-run of the linked benchmark "
        "using a singly-linked std::forward_list to match the Python implementation's representation, which "
        "reproduced the same pattern of ratios rather than narrowing it. The findings support teaching Big-O together "
        "with implementation-aware measurement, while the single-machine scope, absence of Java and hardware-counter "
        "data, and submicrosecond C++ measurements make the work a pilot rather than a final general-purpose "
        "benchmark."
    )
    para(doc, abstract, style="03_Abstract")
    para(doc, "Keywords—data structures; Big-O; empirical algorithmics; Python; C++; microbenchmarking; reproducibility; construct validity", style="04_Keywords")

    # I. INTRODUCTION
    para(doc, "Introduction", style="Heading 1")
    para(doc, f"Big-O notation provides a machine-independent description of growth as input size increases. It is indispensable for algorithm design, yet it intentionally suppresses constant factors, representation costs and hardware effects [{REF['cormen']}]. Consequently, two implementations with the same asymptotic class can differ substantially in observed runtime, while a theoretically favourable structure may carry a higher construction or traversal cost for a particular workload.", style="Body Text")
    para(doc, f"This distinction is especially important in undergraduate computing education. Students often learn that hash-table lookup is expected O(1), array search is O(n), and linked-list traversal is O(n), but they may not observe how language runtimes, boxed objects, allocation strategies and contiguous memory influence actual measurements. Computer architecture texts likewise emphasise that locality and the memory hierarchy shape performance beyond instruction counts [{REF['hennessy']}].", style="Body Text")
    para(doc, "The objective of this pilot is not to declare one programming language universally superior. It is to test whether a small, reproducible experiment can connect theoretical complexity with measured behaviour while making its assumptions and limitations explicit. The resulting protocol is intended as a foundation for a student research project and a later multi-platform study.", style="Body Text")

    para(doc, "Research Questions", style="Heading 2")
    for text in (
        "RQ1: How do language and data-structure implementation affect observed runtime for equivalent workloads?",
        "RQ2: Do measured scaling patterns agree with the workload-level complexity predicted from Big-O analysis?",
        "RQ3: Which methodological limitations must be addressed before extending the pilot into a journal-ready benchmark?",
    ):
        para(doc, text, style="11_Bullet List")

    para(doc, "Contributions", style="Heading 2")
    for text in (
        "A common, deterministic dataset and workload definition shared by Python and C++ implementations.",
        "A reproducible execution pipeline with correctness checks, raw-data hashing and environment metadata.",
        "An empirical distinction between per-operation complexity and the complexity of a batch whose operation count also grows with n.",
        "A transparent account of measurement-resolution, implementation-equivalence and external-validity limitations.",
        "An empirical test of whether the C++/Python linked-container mismatch inflates the observed language gap, rather than leaving that concern as an unverified caveat.",
    ):
        para(doc, text, style="11_Bullet List")

    # II. BACKGROUND
    para(doc, "Background and Related Work", style="Heading 1")
    para(doc, f"The analysis of algorithms separates growth rate from implementation detail [{REF['cormen']}]. Experimental algorithmics complements that abstraction by examining implementations on specified workloads and machines. Algorithm engineering has been described as a methodology connecting design, implementation, experimentation and refinement [{REF['sanders']}]. This perspective motivates reporting sufficient detail to reproduce not only the code but also the experimental conditions.", style="Body Text")
    para(doc, f"Performance measurement is vulnerable to effects that are easy to overlook. Mytkowicz et al. showed that apparently harmless environmental changes can alter conclusions [{REF['mytkowicz']}]. Georges et al. and Kalibera and Jones argued for warm-ups, repeated measurements, uncertainty reporting and explicit treatment of runtime-system variability [{REF['georges']},{REF['kalibera']}]. Arcuri and Briand outlined complementary statistical-testing practice for evaluating randomized algorithms in software engineering, including effect-size reporting alongside significance testing [{REF['arcuri']}]. Fleming and Wallace further cautioned that benchmark summaries can mislead when inappropriate averages are used [{REF['fleming']}]. Reproducibility itself depends on preserving code, data and environment details rather than reporting summary numbers alone [{REF['sandve']},{REF['peng']}]. The present pilot therefore prioritises medians, dispersion, warm-ups, immutable raw results and correctness checks.", style="Body Text")
    para(doc, "The study is deliberately modest. It does not claim that its three containers exhaust data-structure design or that Python and C++ expose equivalent internal representations. Instead, it investigates idiomatic implementations under a shared external workload. That decision improves classroom relevance but weakens causal attribution: measured differences reflect the combined effect of language, runtime, library implementation and element representation.", style="Body Text")

    # III. MATERIALS AND METHODS
    para(doc, "Materials and Methods", style="Heading 1")
    para(doc, "Execution Environment", style="Heading 2")
    env_rows = [
        ["Computer", "MacBook Air (Mac14,2)"],
        ["Processor", "Apple M2, 8 cores (4 performance, 4 efficiency), arm64"],
        ["Memory", "8 GB unified memory"],
        ["Operating system", "macOS 26.5.2, Darwin 25.5.0"],
        ["Python", "CPython 3.13.5"],
        ["C++", "C++17, Apple Clang 21.0.0, -O2"],
        ["Java", "Not executed: no working JDK was available"],
        ["Hardware counters", "Not collected: Linux perf unavailable on macOS"],
        ["Execution date", "23 August 2026"],
    ]
    add_table(doc, "Recorded execution environment.", ["Item", "Recorded value"], env_rows, [1.1, 2.3])

    para(doc, "Structures and Workload", style="Heading 2")
    para(doc, f"The experiment used three abstract structures. The dynamic-array condition used Python list [{REF['pythondocs']}] and C++ std::vector<int> [{REF['isocpp']}]; the linked condition used a custom Python singly linked list and C++ std::list<int>; and the hash condition used Python dict and C++ std::unordered_map<int,int>. Inputs contained unique integers in a deterministic shuffled order.", style="Body Text")
    method_rows = [
        ["Input size", "1,000; 5,000; 10,000; 25,000"],
        ["Build", "Construct the structure from all n values"],
        ["Search", "Query max(10, 0.01n) keys; half present, half absent"],
        ["Delete", "Delete max(1, 0.01n) keys known to be present"],
        ["Traverse", "Sum all keys remaining after deletion"],
        ["Correctness", "Compare search-hit count and post-deletion checksum"],
        ["Timing boundary", "Exclude dataset parsing and CSV output"],
    ]
    add_table(doc, "Common experimental workload.", ["Component", "Definition"], method_rows, [1.0, 2.4])

    para(doc, "Dataset Generation and Validation", style="Heading 2")
    para(doc, "A fixed seed (20260823) generated one text dataset per input size. Every language read the same files. A manifest stored each file's SHA-256 digest. Before performance analysis, all result groups were checked for identical search-hit counts and identical post-deletion checksums. Any disagreement would have invalidated the performance data; none occurred.", style="Body Text")

    para(doc, "Measurement Protocol", style="Heading 2")
    para(doc, "For every language-structure-size combination, the program performed three unrecorded warm-ups followed by ten recorded repetitions. Jobs were shuffled with the experiment seed. Operation durations were measured using monotonic high-resolution clocks and reported in nanoseconds. The complete design produced 2 languages x 3 structures x 4 sizes x 10 repetitions = 240 records.", style="Body Text")
    p = para(doc, style="Body Text")
    p.add_run("Measurement caveat: ").bold = True
    p.add_run("Some optimized C++ workloads completed in less than one microsecond, and the C++ hash-search median at n = 1,000 was recorded as zero. Such values are below a dependable single-shot timing range and are treated as resolution-limited rather than literal zero-cost operations.")

    para(doc, "Analysis", style="Heading 2")
    para(doc, f"The primary summary was the median across ten repetitions, with the mean, standard deviation and non-parametric bootstrap interval retained in the generated analysis file [{REF['bootstrap']}]. Coefficient of variation (CV) was used to describe group dispersion. Python-to-C++ median ratios at n = 25,000 were calculated descriptively; Cliff's delta [{REF['cliff']}], following the A-measure formulation of Vargha and Delaney [{REF['varghadelaney']}], was also computed. Scaling exponents between n = 5,000 and n = 25,000 were estimated as log(T2/T1) / log(n2/n1). No causal language claim is made because language and container representation were not independently manipulated.", style="Body Text")

    # IV. RESULTS
    para(doc, "Results", style="Heading 1")
    para(doc, "Correctness and Measurement Stability", style="Heading 2")
    para(doc, f"All 240 records passed the hit-count and checksum validation. The median CV across the 96 language-structure-size-operation groups was {METRICS['median_group_cv']*100:.1f}%. The maximum CV was {METRICS['max_group_cv']*100:.1f}% and occurred in the resolution-limited C++ hash-search group at n = 1,000. The data therefore show generally stable repeated measurements while also identifying the exact region in which the timer cannot support strong absolute claims.", style="Body Text")

    para(doc, "Search Scaling", style="Heading 2")
    para(doc, "Sequential search rose much more sharply than hash search. Between 5,000 and 25,000 elements, the estimated search exponent was 1.99 for the C++ array and 2.00 for the Python array; the linked implementations produced 1.91 and 1.99, respectively. The corresponding hash exponents were 0.90 for C++ and 0.97 for Python. This apparent quadratic-versus-linear contrast is consistent with the workload definition: the number of search queries itself increases in proportion to n, so an O(n) sequential search repeated O(n) times yields an O(n^2) batch, whereas expected O(1) hash lookup repeated O(n) times yields an O(n) batch.", style="Body Text")

    new_section(doc, cols=1)
    add_figure(doc, GENERATED / "figure_1_search_scaling.png", "Median search-workload time across input sizes. The zero C++ hash-search median at n = 1,000 is omitted from the logarithmic plot.", 6.8)
    new_section(doc, cols=2)

    para(doc, "Runtime at the Largest Input", style="Heading 2")
    new_section(doc, cols=1)
    result_rows = []
    for structure in ("array", "linked", "hash"):
        for operation in ("insert", "search", "delete", "traverse"):
            entry = METRICS["n25000"][structure][operation]
            result_rows.append([
                structure.title(), operation.title(), fmt_ms(entry["cpp"]["median_ms"]),
                fmt_ms(entry["python"]["median_ms"]), f"{entry['python_to_cpp_median_ratio']:.2f}x",
            ])
    add_table(doc, "Median workload runtime at n = 25,000 (ten repetitions).", ["Structure", "Operation", "C++17", "Python 3.13", "Python/C++"], result_rows, [1.0, 1.0, 1.2, 1.2, 1.1])
    new_section(doc, cols=2)
    ratio_min, ratio_max = METRICS["language_ratio_range_n25000"]
    para(doc, f"At n = 25,000, every Python median exceeded its corresponding C++ median; the descriptive ratios ranged from {ratio_min:.2f}x for hash deletion to {ratio_max:.2f}x for array traversal. Cliff's delta was 1.00 for all twelve comparisons, meaning every recorded Python duration in each comparison exceeded every corresponding C++ duration. These values characterise the tested implementations on this Mac and must not be interpreted as universal language speed factors.", style="Body Text")

    new_section(doc, cols=1)
    add_figure(doc, GENERATED / "figure_2_language_ratios.png", "Python-to-C++ median runtime ratios at n = 25,000. The horizontal axis is logarithmic.", 6.8)
    new_section(doc, cols=2)

    para(doc, "Deletion and Traversal Patterns", style="Heading 2")
    para(doc, "Array and linked deletion workloads also approached quadratic scaling because 1% of n keys were deleted and each deletion required sequential location; estimated exponents ranged from 1.85 to 1.89 for linked deletion and 1.87 to 1.89 for array deletion. Hash deletion remained close to linear at 1.00 in C++ and 1.02 in Python. Traversal generally approached linear growth, with exponents between 0.84 and 0.97 across the structures and languages. These patterns are more defensible than the smallest absolute timings because they rely on changes across larger workloads.", style="Body Text")

    para(doc, "Construct-Validity Check: A Singly-Linked C++ Comparison", style="Heading 2")
    para(doc, "Section VI-B notes a representational mismatch in the linked condition: the C++ side used std::list, which is typically doubly linked, while the Python side used a custom singly linked list. To test whether this mismatch inflates the reported ratios, the C++ linked benchmark was re-run using std::forward_list, a singly-linked container matching the Python implementation's structure, across all four input sizes with the same protocol. Every forward_list run reproduced the same search-hit count and post-deletion checksum as both the original std::list run and the Python run at every input size.", style="Body Text")
    fwd = METRICS["linked_fwd_supplement"]["n25000"]
    fwd_rows = []
    for operation in ("insert", "search", "delete", "traverse"):
        entry = fwd[operation]
        fwd_rows.append([
            operation.title(), fmt_ms(entry["cpp_list_median_ms"]), fmt_ms(entry["cpp_forward_list_median_ms"]),
            fmt_ms(entry["python_median_ms"]), f"{entry['python_to_forward_list_ratio']:.2f}x",
        ])
    new_section(doc, cols=1)
    add_table(doc, "Construct-validity check: median linked-structure runtime at n = 25,000 using a singly-linked C++ container.", ["Operation", "C++ list", "C++ fwd_list", "Python 3.13", "Py/fwd_list"], fwd_rows, [0.9, 1.0, 1.1, 1.1, 1.1])
    new_section(doc, cols=2)
    para(doc, f"Matching the C++ container's linkage discipline to the Python implementation did not narrow the measured language gap. The Python-to-C++ ratio was higher under the forward_list comparison than under std::list for search ({fwd['search']['python_to_forward_list_ratio']:.2f}x versus {fwd['search']['python_to_list_ratio']:.2f}x), deletion ({fwd['delete']['python_to_forward_list_ratio']:.2f}x versus {fwd['delete']['python_to_list_ratio']:.2f}x) and traversal ({fwd['traverse']['python_to_forward_list_ratio']:.2f}x versus {fwd['traverse']['python_to_list_ratio']:.2f}x), and only slightly lower for insertion ({fwd['insert']['python_to_forward_list_ratio']:.2f}x versus {fwd['insert']['python_to_list_ratio']:.2f}x). The construct-validity concern therefore does not appear to be inflating the reported Python/C++ ratios for the linked condition.", style="Body Text")

    # V. DISCUSSION
    para(doc, "Discussion", style="Heading 1")
    para(doc, "Big-O Explained, Not Contradicted", style="Heading 2")
    para(doc, "The results do not show a failure of asymptotic analysis. Instead, they demonstrate why the unit of analysis must be stated carefully. A single sequential search is O(n), but this experiment schedules 0.01n searches, creating an O(n^2) batch. A hash lookup is expected O(1) on average, and the corresponding O(n) batch grew close to linearly. The empirical exponents therefore connect measured behaviour to the composed workload rather than replacing theory.", style="Body Text")

    para(doc, "Language and Representation Effects", style="Heading 2")
    para(doc, f"The Python/C++ differences include many inseparable factors: interpretation versus optimized native code, boxed Python integers versus C++ primitive integers, object allocation, iterator implementation and library design. The largest ratio occurred for contiguous traversal, where optimized native loops can exploit compact primitive storage, spatial locality and compiler transformations that pointer-chasing and boxed representations largely defeat [{REF['chilimbi']},{REF['drepper']}]. Hash insertion and deletion showed smaller ratios, indicating that language overhead is not a single constant applied uniformly to every structure and operation.", style="Body Text")

    para(doc, "Educational Value", style="Heading 2")
    para(doc, f"A useful teaching sequence is therefore: predict complexity, define the entire workload, implement correctness checks, measure, and then explain both agreement and disagreement. Students should learn to ask not only 'What is the Big-O?' but also 'What exactly is repeated, what representation is used, what is inside the timed region, and is the measurement long enough for the clock?' This echoes earlier calls in computing education to pair algorithmic analysis with empirical measurement rather than treating asymptotic notation as a substitute for it [{REF['astrachan']}]. The repository makes those questions visible and reproducible.", style="Body Text")

    # VI. THREATS TO VALIDITY
    para(doc, "Threats to Validity", style="Heading 1")
    para(doc, "Internal Validity", style="Heading 2")
    para(doc, "Background activity, thermal conditions and operating-system scheduling were not instrumented. Although jobs were shuffled and measurements were repeated, the study did not pin processes to performance cores. Extremely short C++ operations are vulnerable to timer quantisation and clock-call overhead. A revised benchmark should batch pure operations until each timed interval exceeds a predefined duration and should separately quantify timing overhead.", style="Body Text")

    para(doc, "Construct Validity", style="Heading 2")
    para(doc, "The structures are functionally comparable but not internally identical. Python's custom singly linked list differs from C++ std::list, which is typically doubly linked. Python list holds object references, whereas std::vector<int> stores primitive integers contiguously. Consequently, the experiment evaluates idiomatic implementation conditions rather than isolating a pure language effect. This concern is tested empirically in Section IV-E, which reruns the C++ linked benchmark with the singly-linked std::forward_list; the resulting ratios do not shrink under the more closely matched comparison. Peak memory and cache behaviour were proposed in the broader project but were not measured on this Mac and are not inferred in this paper.", style="Body Text")

    para(doc, "External and Conclusion Validity", style="Heading 2")
    para(doc, "One Apple M2 laptop, one Python version and one C++ toolchain cannot represent other processors, operating systems, compilers or runtime configurations. Only four sizes and one random-input family were examined. Ten repetitions support descriptive stability but not broad population inference. Java was not executed because no working JDK was available. The conclusions are therefore restricted to this environment and protocol.", style="Body Text")

    # VII. CONCLUSION
    para(doc, "Conclusion and Future Work", style="Heading 1")
    para(doc, "This reproducible Mac pilot showed that theoretical growth and observed runtime can be taught together. Sequential search and deletion batches grew close to quadratically when both collection size and operation count increased, while hash batches grew close to linearly. Python and C++ showed substantial but operation-dependent runtime differences, and all correctness outputs agreed. The study also exposed a measurement weakness: several optimized C++ workloads were too short for reliable single-shot timing.", style="Body Text")
    para(doc, "Before journal submission, the benchmark should use calibrated operation batching, add Java 17, collect peak resident memory and hardware cache counters on a controlled Linux machine, record temperature and power mode, test additional input distributions, and reproduce the experiment on at least one independent system. Those extensions should be reported as new evidence, not combined silently with the present Mac data.", style="Body Text")

    # Component heads (unnumbered)
    para(doc, "Acknowledgment", style="Heading 5")
    para(doc, "Not applicable.", style="Body Text")

    para(doc, "Declarations", style="Heading 5")
    p = para(doc, style="Body Text")
    p.add_run("Conflict of Interest: ").bold = True
    p.add_run("The authors declare no conflict of interest.")
    p = para(doc, style="Body Text")
    p.add_run("Funding: ").bold = True
    p.add_run("This research received no external funding.")
    p = para(doc, style="Body Text")
    p.add_run("Data and Code Availability: ").bold = True
    p.add_run(
        f"The repository contains the dataset generator, Python and C++ implementations, Java source, validation "
        f"tests, execution configuration, raw CSV records and analysis scripts. The raw combined CSV contains "
        f"{METRICS['record_count']} records (SHA-256: {METRICS['raw_sha256']}); the supplementary construct-validity "
        f"check adds {METRICS['linked_fwd_supplement']['record_count']} further records (SHA-256: "
        f"{METRICS['linked_fwd_supplement']['raw_sha256']}). The repository is publicly available at "
        f"https://github.com/maheshchudaman/beyond-big-o-research and is archived on Zenodo with concept DOI "
        f"10.5281/zenodo.22089928."
    )
    p = para(doc, style="Body Text")
    p.add_run("Generative AI Use: ").bold = True
    p.add_run(
        "Generative AI (Codex/Claude) assisted with repository scaffolding, code review, execution orchestration, "
        "analysis scripting and manuscript drafting. All numerical claims were generated programmatically from the "
        "preserved raw CSV; no measurements were fabricated. The named authors independently verified the code, "
        "results, citations and wording, and take full responsibility for the content."
    )

    para(doc, "References", style="Heading 5")
    for _, reference in REFERENCES:
        para(doc, reference, style="17_References")

    doc.core_properties.title = "Beyond Big-O: A Reproducible Apple-Silicon Pilot Study of Runtime Performance for Common Data Structures in Python and C++"
    doc.core_properties.author = "Mahesh Patil"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build()
