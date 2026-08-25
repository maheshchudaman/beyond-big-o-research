#!/usr/bin/env python3
"""Build the Beyond Big-O student research workbook."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "workbook" / "Beyond_Big_O_Student_Research_Workbook.docx"
BLUE = "2E74B5"
DARK_BLUE = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GOLD = "C38D2E"
MUTED = "5E6B78"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=11, bold=False, color=DARK_BLUE, italic=False, font="Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, *, size=11, bold=False, italic=False, color=DARK_BLUE, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_run(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_run(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_callout(doc, label: str, text: str, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{label}: "), bold=True, color=DARK_BLUE)
    set_run(p.add_run(text), color=DARK_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), size=9.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if index == 0 and len(value) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(value), size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def blank_lines(doc, count=3):
    for _ in range(count):
        p = doc.add_paragraph("________________________________________________________________________________")
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            set_run(run, color="AAB4BE", size=9)


def page_break(doc):
    doc.add_page_break()


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("BEYOND BIG-O  |  STUDENT RESEARCH WORKBOOK"), size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("Reproducible research • Version 0.1 • August 2026"), size=8.5, color=MUTED)


def build():
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc.sections[0])

    # Cover — editorial_cover override.
    add_text(doc, "RESEARCH WORKBOOK", size=11, bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, before=72, after=18)
    add_text(doc, "BEYOND BIG-O", size=31, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_text(doc, "A Cross-Language Empirical Study of Runtime, Memory and Cache Performance of Common Data Structures", size=15, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    add_text(doc, "Python  •  Java  •  C++", size=12, bold=True, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, after=54)
    add_callout(doc, "Purpose", "Guide third-year Computer Engineering students from a research question to a reproducible experiment, validated code, defensible analysis and an ethical journal manuscript.")
    add_text(doc, "Project mentor: Mahesh Patil", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=42, after=4)
    add_text(doc, "First research-framework edition • August 2026", size=9.5, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break(doc)
    add_heading(doc, "How to Use This Workbook", 1)
    add_text(doc, "Complete the workbook alongside the GitHub repository. Do not move to a later gate until the preceding gate has been reviewed and signed off.")
    add_table(doc, ["Gate", "Evidence required", "Status"], [
        ["1", "Research questions, hypotheses and scope approved", "☐"],
        ["2", "Common datasets generated; manifest preserved", "☐"],
        ["3", "All implementations pass correctness checks", "☐"],
        ["4", "Pilot experiment reviewed; timing protocol frozen", "☐"],
        ["5", "Full experiment completed without selective deletion", "☐"],
        ["6", "Statistical analysis and figures independently checked", "☐"],
        ["7", "Paper, code and research-integrity checklist complete", "☐"],
    ], [900, 7200, 1260])
    add_heading(doc, "Team Agreement", 2)
    add_bullet(doc, "Every student must understand at least one implementation and the complete experimental protocol.")
    add_bullet(doc, "No measurement may be edited, fabricated or omitted merely because it conflicts with expectations.")
    add_bullet(doc, "AI assistance, external code and datasets must be disclosed according to institutional and journal policy.")
    add_bullet(doc, "Authorship must reflect substantial intellectual and practical contributions.")
    add_text(doc, "Team signatures and date:", bold=True)
    blank_lines(doc, 3)

    page_break(doc)
    add_heading(doc, "1. Research Identity and Team Roles", 1)
    add_table(doc, ["Field", "Entry"], [
        ["Institution / department", ""], ["Faculty mentor", ""], ["Student team", ""],
        ["Project start date", ""], ["Target completion date", ""], ["GitHub repository URL", ""],
    ], [2400, 6960])
    add_heading(doc, "Role allocation", 2)
    add_table(doc, ["Role", "Primary owner", "Reviewer", "Deliverable"], [
        ["Research lead", "", "", "Protocol and decision log"],
        ["Python lead", "", "", "Implementation and tests"],
        ["Java lead", "", "", "Implementation and JVM notes"],
        ["C++ lead", "", "", "Implementation and compiler notes"],
        ["Data/analysis lead", "", "", "Validated datasets and analysis"],
        ["Reproducibility lead", "", "", "Environment log and rerun"],
    ], [1900, 1800, 1800, 3860])
    add_callout(doc, "Rule", "A primary owner may write the code, but a different student must review its correctness before benchmarking.", fill=LIGHT_GRAY)

    page_break(doc)
    add_heading(doc, "2. Research Questions and Hypotheses", 1)
    add_heading(doc, "Approved research questions", 2)
    for text in (
        "RQ1. How does language choice change the observed cost of equivalent data-structure operations?",
        "RQ2. At what input sizes do theoretical advantages become observable in practice?",
        "RQ3. How do access pattern and deletion rate affect runtime and cache behaviour?",
        "RQ4. Can a normalised suitability score recommend a structure for different resource priorities?",
    ):
        add_bullet(doc, text)
    add_heading(doc, "Hypothesis worksheet", 2)
    add_table(doc, ["ID", "Prediction", "Reasoning", "Evidence that could reject it"], [
        ["H1", "Hash search becomes faster at larger n.", "", ""],
        ["H2", "Linked structures show poorer cache locality.", "", ""],
        ["H3", "Language and structure interact significantly.", "", ""],
    ], [800, 2500, 2860, 3200])
    add_heading(doc, "Novel contribution", 2)
    add_text(doc, "State the contribution in one sentence. Avoid claiming that a comparison alone is novel.")
    blank_lines(doc, 3)

    page_break(doc)
    add_heading(doc, "3. Experimental Design", 1)
    add_table(doc, ["Variable type", "Variables"], [
        ["Independent", "Language, structure, input size and workload"],
        ["Dependent", "Operation time, peak RSS, cache references, cache misses, correctness"],
        ["Controlled", "Machine, dataset, seed, compiler flags, warm-ups, power mode"],
        ["Blocking / nuisance", "Execution order, background load, temperature, garbage collection"],
    ], [2200, 7160])
    add_heading(doc, "Workload matrix", 2)
    add_table(doc, ["Operation", "Exact definition", "Included in timed region?", "Validation"], [
        ["Build", "Construct from all input values", "Yes", "Final element count"],
        ["Search", "Run the same present and absent queries", "Yes", "Hit count"],
        ["Delete", "Delete the same present keys", "Yes", "Remaining checksum"],
        ["Traverse", "Sum all remaining keys", "Yes", "Checksum"],
        ["Dataset loading", "Read and parse the file", "No", "SHA-256 manifest"],
        ["CSV writing", "Persist the result row", "No", "Schema check"],
    ], [1500, 3900, 1800, 2160])
    add_callout(doc, "Critical distinction", "The experiment compares idiomatic library structures, not identical internal implementations. Report this as a construct-validity limitation.")

    page_break(doc)
    add_heading(doc, "4. Environment and Reproducibility Log", 1)
    add_table(doc, ["Item", "Recorded value"], [
        ["CPU model / cores", ""], ["RAM capacity", ""], ["Operating system / kernel", ""],
        ["Python version", ""], ["Java / JVM version", ""], ["C++ compiler and version", ""],
        ["Compiler flags", "-O2 -std=c++17"], ["Power mode", ""], ["Background services controlled", ""],
        ["Linux perf permission status", ""], ["Dataset manifest hash", ""], ["Git commit hash", ""],
    ], [3100, 6260])
    add_heading(doc, "Pre-run checklist", 2)
    for item in (
        "☐ Machine connected to stable power and unnecessary applications closed",
        "☐ Dataset manifest verified and repository working tree clean",
        "☐ Correct release build used; debug build excluded",
        "☐ Warm-up count and repetition count match the frozen configuration",
        "☐ System clock, thermal state and power mode recorded",
    ):
        add_text(doc, item, after=4)

    page_break(doc)
    add_heading(doc, "5. GitHub and Code-Review Gate", 1)
    add_heading(doc, "Repository checklist", 2)
    for item in (
        "☐ Professional GitHub accounts created; two-factor authentication enabled",
        "☐ Main branch protected; work completed through reviewed pull requests",
        "☐ README explains setup, fairness rules and repository map",
        "☐ No credentials, generated binaries or manually edited measurements committed",
        "☐ GitHub Actions validation passes on the exact commit used for experiments",
    ):
        add_text(doc, item, after=4)
    add_heading(doc, "Implementation review record", 2)
    add_table(doc, ["Language / structure", "Correctness reviewer", "Test evidence", "Approved"], [
        ["Python / array", "", "", "☐"], ["Python / linked", "", "", "☐"], ["Python / hash", "", "", "☐"],
        ["Java / array", "", "", "☐"], ["Java / linked", "", "", "☐"], ["Java / hash", "", "", "☐"],
        ["C++ / array", "", "", "☐"], ["C++ / linked", "", "", "☐"], ["C++ / hash", "", "", "☐"],
    ], [2100, 2300, 3600, 1360])

    page_break(doc)
    add_heading(doc, "6. Pilot Experiment and Validation", 1)
    add_number(doc, "Generate the smallest dataset and preserve its manifest.")
    add_number(doc, "Run all structures in each available language with one recorded repetition.")
    add_number(doc, "Confirm that query hits and post-deletion checksums agree.")
    add_number(doc, "Inspect unusually small, zero or unstable timing values.")
    add_number(doc, "Repeat the pilot after a clean rebuild.")
    add_number(doc, "Freeze the protocol only after the pilot passes.")
    add_heading(doc, "Pilot outcome", 2)
    add_table(doc, ["Check", "Expected", "Observed", "Decision"], [
        ["Hit counts", "Identical", "", ""], ["Checksums", "Identical", "", ""],
        ["Recorded rows", "Complete", "", ""], ["Timing variability", "Explainable", "", ""],
        ["Environment metadata", "Complete", "", ""],
    ], [2200, 1900, 2700, 2560])
    add_text(doc, "Protocol changes made after pilot:", bold=True)
    blank_lines(doc, 4)

    page_break(doc)
    add_heading(doc, "7. Full Experiment Run Log", 1)
    add_callout(doc, "Stop condition", "Stop the experiment immediately if correctness values disagree. Performance data from an incorrect implementation is invalid.", fill="FCE8E6")
    add_table(doc, ["Run date/time", "Commit", "Dataset(s)", "Execution order seed", "Status / notes"], [
        ["", "", "", "", ""], ["", "", "", "", ""], ["", "", "", "", ""],
        ["", "", "", "", ""], ["", "", "", "", ""], ["", "", "", "", ""],
        ["", "", "", "", ""], ["", "", "", "", "",],
    ], [1600, 1300, 1800, 1900, 2760])
    add_heading(doc, "Deviation log", 2)
    add_text(doc, "Record every interruption, failed build, thermal concern, permission problem or rerun. Do not silently replace a run.")
    blank_lines(doc, 6)

    page_break(doc)
    add_heading(doc, "8. Analysis Plan", 1)
    add_table(doc, ["Question", "Required analysis", "Minimum visual"], [
        ["RQ1", "Language × structure comparison by operation", "Distribution plus median trend"],
        ["RQ2", "Interaction with input size on logarithmic axes", "Scaling curve"],
        ["RQ3", "Workload and cache-miss comparison", "Effect plot"],
        ["RQ4", "DSSS sensitivity across weight combinations", "Ranking stability plot"],
    ], [1100, 5200, 3060])
    add_heading(doc, "Statistical checklist", 2)
    for item in (
        "☐ Report sample count, mean, median, standard deviation and 95% confidence interval",
        "☐ Examine distributions and outliers before selecting tests",
        "☐ Report effect sizes with p-values",
        "☐ Correct for multiple comparisons",
        "☐ Separate statistical significance from practical significance",
        "☐ Include negative and unexpected findings",
        "☐ Keep raw data immutable and analysis code versioned",
    ):
        add_text(doc, item, after=4)
    add_text(doc, "Primary statistical test and justification:", bold=True)
    blank_lines(doc, 3)

    page_break(doc)
    add_heading(doc, "9. Data Structure Suitability Score", 1)
    add_text(doc, "For structure s, calculate:", bold=True)
    add_text(doc, "DSSS(s) = wₜ·Tₙₒᵣₘ(s) + wₘ·Mₙₒᵣₘ(s) + w꜀·Cₙₒᵣₘ(s)", size=15, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=14)
    add_table(doc, ["Term", "Meaning"], [
        ["Tₙₒᵣₘ", "Normalised execution-time measure"],
        ["Mₙₒᵣₘ", "Normalised memory measure"],
        ["Cₙₒᵣₘ", "Normalised cache-miss measure"],
        ["wₜ, wₘ, w꜀", "Application-priority weights; together they must equal 1"],
    ], [2200, 7160])
    add_heading(doc, "Mandatory sensitivity cases", 2)
    add_table(doc, ["Scenario", "Time weight", "Memory weight", "Cache weight", "Expected use"], [
        ["Balanced", "0.34", "0.33", "0.33", "General-purpose"],
        ["Speed-first", "0.70", "0.15", "0.15", "Latency-sensitive"],
        ["Memory-first", "0.15", "0.70", "0.15", "Memory-limited"],
        ["Cache-first", "0.15", "0.15", "0.70", "Locality-sensitive"],
    ], [1800, 1500, 1700, 1600, 2760])
    add_callout(doc, "Interpretation", "DSSS must support—not replace—the raw measurements. A ranking without sensitivity analysis would be arbitrary.")

    page_break(doc)
    add_heading(doc, "10. Paper Blueprint", 1)
    add_table(doc, ["Section", "Required content", "Owner", "Ready"], [
        ["Abstract", "Problem, method, principal result, significance", "", "☐"],
        ["Introduction", "Motivation, gap, questions and contribution", "", "☐"],
        ["Related work", "Critical comparison; no catalogue of summaries", "", "☐"],
        ["Method", "Data, structures, operations, controls and statistics", "", "☐"],
        ["Results", "Validated measurements, uncertainty and figures", "", "☐"],
        ["Discussion", "Meaning, theory–practice differences and trade-offs", "", "☐"],
        ["Validity", "Internal, construct, external and conclusion validity", "", "☐"],
        ["Conclusion", "Direct answers, limitations and next study", "", "☐"],
        ["Availability", "Repository, commit, data and reproduction command", "", "☐"],
    ], [1500, 5600, 1260, 1000])
    add_heading(doc, "One-sentence contribution test", 2)
    add_text(doc, "Complete: This study contributes ________________________________________________")
    blank_lines(doc, 3)

    page_break(doc)
    add_heading(doc, "11. Research Integrity and Journal Check", 1)
    add_heading(doc, "Integrity declaration", 2)
    for item in (
        "☐ All reported measurements were produced by the disclosed code and hardware",
        "☐ No inconvenient result was removed without a documented, protocol-based reason",
        "☐ External code, datasets and AI assistance were disclosed",
        "☐ Every author reviewed the final manuscript and underlying evidence",
        "☐ The repository commit corresponding to the paper is preserved",
    ):
        add_text(doc, item, after=4)
    add_heading(doc, "Journal screening", 2)
    add_table(doc, ["Check", "Evidence"], [
        ["Journal scope matches empirical software/systems research", ""],
        ["Current Scopus/SJR quartile verified for the relevant category and year", ""],
        ["Publisher and editorial board are transparent", ""],
        ["Peer-review process and fees are clearly stated", ""],
        ["No guaranteed acceptance or suspiciously rapid publication promise", ""],
        ["Author guidelines and research-data policy reviewed", ""],
    ], [5000, 4360])
    add_callout(doc, "Reminder", "Q4 is a ranking category, not permission to submit weak or unreproducible work. Quartiles can change by year and subject category.", fill=LIGHT_GRAY)

    page_break(doc)
    add_heading(doc, "12. Twelve-Week Execution Plan", 1)
    add_table(doc, ["Week", "Focus", "Exit evidence"], [
        ["1", "Accounts, Git and topic orientation", "Repository and team roles"],
        ["2", "Focused literature search", "Evidence matrix and gap statement"],
        ["3", "Research questions and preregistration", "Frozen draft protocol"],
        ["4", "Dataset generator and manifest", "Validated common inputs"],
        ["5", "Python implementation", "Tests and review"],
        ["6", "C++ implementation", "Tests and review"],
        ["7", "Java implementation", "Tests and review"],
        ["8", "Pilot and protocol correction", "Pilot sign-off"],
        ["9", "Full experiments", "Immutable raw data"],
        ["10", "Statistics and visualisation", "Checked analysis outputs"],
        ["11", "Manuscript drafting", "Complete internal draft"],
        ["12", "Independent reproduction and revision", "Submission-ready package"],
    ], [900, 3900, 4560])
    add_heading(doc, "Weekly mentor review", 2)
    add_table(doc, ["Date", "Evidence reviewed", "Decision / action", "Initials"], [["", "", "", ""] for _ in range(5)], [1400, 3400, 3400, 1160])

    page_break(doc)
    add_heading(doc, "Appendix A. Essential Commands", 1)
    commands = [
        ("Generate datasets", "python3 scripts/generate_datasets.py"),
        ("Run unit tests", "python3 -m unittest discover -s tests -v"),
        ("Run smoke test", "python3 scripts/run_all.py --smoke"),
        ("Run full benchmark", "python3 scripts/run_all.py"),
        ("Analyse results", "python3 scripts/analyse_results.py"),
        ("Check repository", "git status"),
        ("Record exact commit", "git rev-parse HEAD"),
    ]
    add_table(doc, ["Purpose", "Command"], [[a, b] for a, b in commands], [2500, 6860])
    add_heading(doc, "Linux cache profiling", 2)
    add_text(doc, "Use scripts/profile_linux.sh to capture cache references, cache misses, cycles, instructions and peak resident memory. Hardware counters may require administrator configuration; record any limitation.")
    add_heading(doc, "Final handover package", 2)
    for item in (
        "☐ Manuscript and cover letter",
        "☐ Source code at a tagged commit",
        "☐ Dataset manifest and raw results",
        "☐ Analysis code and processed tables",
        "☐ Environment record and reproduction instructions",
        "☐ Contributor and AI-use declarations",
    ):
        add_text(doc, item, after=4)

    doc.core_properties.title = "Beyond Big-O Student Research Workbook"
    doc.core_properties.subject = "Cross-language data-structure performance research"
    doc.core_properties.author = "Mahesh Patil"
    doc.core_properties.keywords = "Big-O, data structures, Python, Java, C++, cache, reproducibility"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
